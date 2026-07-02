import shutil
import sys
from pathlib import Path

from docx import Document

SOURCE_DOCX = Path(
    r"c:\Users\asadr\OneDrive\Desktop\Asad-ur-Rahman-Resume\Asad-ur-Rahman-Resume-Cybergen.docx"
)
PUBLIC_DIR = Path(__file__).resolve().parents[1] / "public"
PORTFOLIO_URL = "https://portfolio-seven-fawn-68.vercel.app/"
PDF_NAME = "Asad-ur-Rahman-Resume.pdf"
OLD_PDF = PUBLIC_DIR / "AsadResume.pdf"


def update_resume(doc: Document) -> None:
    contact = doc.paragraphs[2]
    if PORTFOLIO_URL not in contact.text:
        contact.text = contact.text.rstrip() + f" | {PORTFOLIO_URL}"

    doc.paragraphs[37].text = (
        "Deployed on Vercel using React, TypeScript, and Tailwind CSS with a modern 3D UI."
    )

    has_live_link = any("Live Portfolio" in paragraph.text for paragraph in doc.paragraphs)
    if not has_live_link:
        doc.paragraphs[38].insert_paragraph_before(f"Live Portfolio: {PORTFOLIO_URL}")


def convert_docx_to_pdf(docx_path: Path, pdf_path: Path) -> None:
    try:
        from docx2pdf import convert

        convert(str(docx_path), str(pdf_path))
        return
    except Exception as error:
        print(f"docx2pdf failed: {error}", file=sys.stderr)

    try:
        import win32com.client

        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(str(docx_path.resolve()))
        doc.SaveAs(str(pdf_path.resolve()), FileFormat=17)
        doc.Close()
        word.Quit()
        return
    except Exception as error:
        raise RuntimeError(f"Could not convert DOCX to PDF: {error}") from error


def main() -> None:
    PUBLIC_DIR.mkdir(parents=True, exist_ok=True)

    output_docx = PUBLIC_DIR / "Asad-ur-Rahman-Resume.docx"
    output_pdf = PUBLIC_DIR / PDF_NAME

    shutil.copy2(SOURCE_DOCX, output_docx)

    doc = Document(output_docx)
    update_resume(doc)
    doc.save(output_docx)

    convert_docx_to_pdf(output_docx, output_pdf)

    if OLD_PDF.exists():
        OLD_PDF.unlink()

    print(f"Saved: {output_docx}")
    print(f"Saved: {output_pdf}")


if __name__ == "__main__":
    main()
